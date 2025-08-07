import pandas as pd
from io import BytesIO
import docx

def read_uploaded_file(uploaded_file):
    """
    Reads uploaded CSV, XLSX, or DOCX file and returns a pandas DataFrame
    with lowercase column names for consistency.
    """

    if uploaded_file is None:
        return None

    filename = uploaded_file.name.lower()

    if filename.endswith(".csv"):
        df = pd.read_csv(uploaded_file)

    elif filename.endswith(".xlsx"):
        df = pd.read_excel(uploaded_file)

    elif filename.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        data = []
        keys = []

        # Attempt to read the first table
        if not doc.tables:
            raise ValueError("DOCX file does not contain any tables.")

        table = doc.tables[0]
        for i, row in enumerate(table.rows):
            text = [cell.text.strip() for cell in row.cells]
            if i == 0:
                keys = text
            else:
                data.append(dict(zip(keys, text)))

        df = pd.DataFrame(data)

    else:
        raise ValueError("Unsupported file format. Only CSV, XLSX, and DOCX are supported.")

    # Normalize column names
    df.columns = df.columns.str.strip().str.lower()
    return df
