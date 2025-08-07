import pandas as pd
import io
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

def generate_report(format: str, df: pd.DataFrame):
    format = format.lower()

    if df.empty:
        raise ValueError("Assignment data is empty.")

    # Ensure all relevant fields are present in the DataFrame
    # Normalize columns to lower case for matching
    df.columns = [c.strip().lower() for c in df.columns]
    # Map possible variations to standard names
    col_map = {}
    for col in df.columns:
        if col in ['student field', 'field', 'department', 'specialization']:
            col_map[col] = 'student field'
        elif col in ['lecturer field', 'field (lecturer)', 'lecturer department', 'lecturer specialization']:
            col_map[col] = 'lecturer field'
        elif col in ['student name', 'name', 'student']:
            col_map[col] = 'student name'
        elif col in ['matric number', 'matric no', 'matric']:
            col_map[col] = 'matric number'
        elif col in ['assigned lecturer', 'lecturer', 'assigned to']:
            col_map[col] = 'assigned lecturer'
        else:
            col_map[col] = col
    df = df.rename(columns=col_map)
    preferred_order = [
        'student name', 'matric number', 'student field',
        'assigned lecturer', 'lecturer field'
    ]
    # Remove 'max_students' if present
    if 'max_students' in df.columns:
        df = df.drop(columns=['max_students'])
    # Add missing columns as empty if not present
    for col in preferred_order:
        if col not in df.columns:
            df[col] = ''
    # Reorder columns for consistency
    ordered_cols = [col for col in preferred_order if col in df.columns] + [col for col in df.columns if col not in preferred_order]
    df = df[ordered_cols]

    if format == "csv":
        buffer = io.StringIO()
        # Write header in uppercase and bold (where supported)
        df.columns = [col.upper() for col in df.columns]
        df.to_csv(buffer, index=False)
        return buffer.getvalue().encode(), "text/csv", "assignment.csv"

    elif format == "excel":
        import openpyxl
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df.columns = [col.upper() for col in df.columns]
            df.to_excel(writer, index=False)
            # Bold header row
            worksheet = writer.sheets['Sheet1']
            for cell in worksheet[1]:
                cell.font = openpyxl.styles.Font(bold=True)
                cell.alignment = openpyxl.styles.Alignment(horizontal='center')
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "assignment.xlsx"

    elif format == "pdf":
        return generate_pdf(df)

    elif format == "word":
        return generate_word(df)

    else:
        raise ValueError(f"Unsupported format: {format}")

# -------------------------
# ✅ PDF Report Generator
# -------------------------
def generate_pdf(df: pd.DataFrame):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    elements.append(Paragraph("Student–Lecturer Assignment Report", styles['Title']))
    elements.append(Spacer(1, 12))

    # Group by lecturer
    group_cols = ["assigned lecturer", "lecturer field"]
    student_cols = ["student name", "matric number", "student field"]
    grouped = df.groupby(group_cols)
    for (lecturer, field), group in grouped:
        heading = f"{lecturer} ({field})" if field else f"{lecturer}"
        elements.append(Paragraph(heading, styles['Heading2']))
        data = [[col.upper() for col in student_cols]] + group[student_cols].values.tolist()
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 18))
    doc.build(elements)
    pdf_data = buffer.getvalue()
    buffer.close()
    return pdf_data, "application/pdf", "assignment.pdf"

# -------------------------
# ✅ Word Report Generator
# -------------------------
def generate_word(df: pd.DataFrame):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("Student–Lecturer Assignment Report", level=1)

    group_cols = ["assigned lecturer", "lecturer field"]
    student_cols = ["student name", "matric number", "student field"]
    grouped = df.groupby(group_cols)
    for (lecturer, field), group in grouped:
        heading = f"{lecturer} ({field})" if field else f"{lecturer}"
        doc.add_heading(heading, level=2)
        table = doc.add_table(rows=1, cols=len(student_cols))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(student_cols):
            hdr_cells[i].text = column.upper()
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = 1  # Center
        for _, row in group.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(student_cols):
                row_cells[i].text = str(row.get(col, ""))
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = 1  # Center
        doc.add_paragraph()
    doc.save(buffer)
    return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "assignment.docx"
