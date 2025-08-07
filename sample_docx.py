from docx import Document
import io

def generate_docx_template_student():
    doc = Document()
    doc.add_heading("Student Template", 0)
    table = doc.add_table(rows=3, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Matric Number'
    hdr_cells[2].text = 'Field'
    row1 = table.rows[1].cells
    row1[0].text = 'John Doe'
    row1[1].text = 'CSC/1234/21'
    row1[2].text = 'Artificial Intelligence'
    row2 = table.rows[2].cells
    row2[0].text = 'Jane Smith'
    row2[1].text = 'CSC/5678/21'
    row2[2].text = 'Artificial Intelligence'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_docx_template_lecturer():
    doc = Document()
    doc.add_heading("Lecturer Template", 0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Max_Students'
    hdr_cells[2].text = 'Field'
    row1 = table.add_row().cells
    row1[0].text = 'Dr. Smith'
    row1[1].text = '8'
    row1[2].text = 'Artificial Intelligence'
    row2 = table.add_row().cells
    row2[0].text = 'Dr. Johnson'
    row2[1].text = '10'
    row2[2].text = 'Artificial Intelligence'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
